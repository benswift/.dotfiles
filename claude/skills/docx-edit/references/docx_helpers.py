"""Reusable helpers for editing a .docx in place with python-docx.

Run with: uv run --with python-docx python your_script.py

The guiding idea: when a .docx is authoritative or hand-edited, you patch it at
the *run* level and verify, rather than regenerating it. These helpers cover the
three things that are fiddly by hand --- rewriting a paragraph's runs while
keeping its formatting, diffing two docx files paragraph-by-paragraph, and
toggling legacy form checkboxes additively.
"""

from __future__ import annotations

import xml.sax.saxutils as su

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn


# --- locating paragraphs -----------------------------------------------------


def find_para(doc, text):
    """Index of the first paragraph whose stripped text == `text` (fail fast).

    Anchor edits on stable heading/landmark text so a structural shift in the
    doc surfaces as a loud failure instead of a silent edit to the wrong place.
    """
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text:
            return i
    raise ValueError(f"paragraph not found: {text!r}")


# --- reading / writing runs --------------------------------------------------


def runs_of(paragraph):
    """[(text, italic, bold), ...] --- the unit you compare and rebuild."""
    return [(r.text, bool(r.italic), bool(r.bold)) for r in paragraph.runs if r.text]


def make_run(text, *, italic=False, bold=False, size_half_pt=None, font_theme=None):
    """Build a <w:r> element carrying its own rPr.

    `xml:space="preserve"` keeps leading/trailing spaces (otherwise Word eats
    them and words run together). `size_half_pt` is Word's half-point unit, so
    10pt == 20. Match `font_theme` / size to the surrounding runs of the doc you
    are editing; inspect an existing run's XML first if unsure.
    """
    rpr = ""
    if font_theme:
        rpr += '<w:rFonts w:asciiTheme="%s" w:hAnsiTheme="%s"/>' % (
            font_theme,
            font_theme,
        )
    if bold:
        rpr += "<w:b/><w:bCs/>"
    if italic:
        rpr += "<w:i/><w:iCs/>"
    if size_half_pt:
        rpr += '<w:sz w:val="%d"/><w:szCs w:val="%d"/>' % (size_half_pt, size_half_pt)
    xml = ('<w:r %s><w:rPr>%s</w:rPr><w:t xml:space="preserve">%s</w:t></w:r>') % (
        nsdecls("w"),
        rpr,
        su.escape(text),
    )
    return parse_xml(xml)


def set_runs(paragraph, runs, **run_kwargs):
    """Replace a paragraph's content with `runs` = [(text, italic, bold), ...].

    Keeps the paragraph's <w:pPr> (style, spacing, numbering) and swaps only the
    run children. `run_kwargs` (size_half_pt, font_theme) pass through to every
    new run. To DROP a sentence, pass the remaining runs; to blank a paragraph,
    pass []. Never rewrite `paragraph.text` directly --- that collapses every run
    into one and loses italics/bold.
    """
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    for text, italic, bold in runs:
        p.append(make_run(text, italic=italic, bold=bold, **run_kwargs))


# --- verifying the edit ------------------------------------------------------


def diff_paragraphs(path_a, path_b):
    """Yield (index, kind, a, b) for every paragraph that differs.

    kind is "count" (whole-doc, paragraph counts diverged), "text", or "format".
    Run this between the pre-edit backup and the patched copy: the only diffs
    should be the paragraphs you meant to touch. Anything else is a bug.
    """
    a = Document(str(path_a)).paragraphs
    b = Document(str(path_b)).paragraphs
    if len(a) != len(b):
        yield (None, "count", len(a), len(b))
        return
    for i, (pa, pb) in enumerate(zip(a, b)):
        ta = "".join(r.text for r in pa.runs)
        tb = "".join(r.text for r in pb.runs)
        if ta != tb:
            yield (i, "text", ta, tb)
        elif runs_of(pa) != runs_of(pb):
            yield (i, "format", runs_of(pa), runs_of(pb))


# --- legacy Word form checkboxes (w:ffData / FORMCHECKBOX) -------------------
#
# These are the grey clickable boxes in protected Word forms, NOT the content
# control or Unicode-glyph kinds. Find the table/cell first, then the checkbox
# elements within it in document order.


def checkboxes_in(cell):
    """All <w:checkBox> elements under a python-docx table cell, in order."""
    return cell._tc.findall(".//" + qn("w:checkBox"))


def is_checked(checkbox):
    """True if ticked. Handles both conventions: a bare <w:checked/> (presence
    == ticked) and a persistent <w:checked w:val="0|1"/> that toggles. Inspect
    your doc once to see which it uses."""
    el = checkbox.find(qn("w:checked"))
    if el is None:
        return False
    return el.get(qn("w:val")) not in ("0", "false")


def set_checked(checkbox, checked=True):
    """Tick/untick additively. Returns True if the state changed.

    Prefer ticking only what you intend and never blanket-clearing: in a shared
    form, an untick you didn't mean to make is a silent data loss.
    """
    el = checkbox.find(qn("w:checked"))
    val = "1" if checked else "0"
    if el is None:
        checkbox.append(parse_xml('<w:checked %s w:val="%s"/>' % (nsdecls("w"), val)))
        return True
    if el.get(qn("w:val")) != val:
        el.set(qn("w:val"), val)
        return True
    return False
